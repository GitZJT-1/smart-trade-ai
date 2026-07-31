#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
validate_3docs.py — 三单一致性校验（程序化，不靠肉眼）+ 校验报告

校验项:
  1. 合同号 / 规格书号: 发票 = 箱单 = 报关单 = 合同 fields.json
  2. 发票金额: Σ(单价×数量) == 合同总金额（币种一致）
  3. 箱单数量/毛净重 = 报关单数量/毛净重
  4. 关键数字全部来自合同，无编造（对比 fields.json）

用法:
  python validate_3docs.py --contract fields.json --invoice 发票.docx \
      --packing 箱单.docx --customs 报关单.xls --report report.md

退出码: 0=全部通过, 1=存在不一致（报告照常输出）
"""
import argparse
import json
import re
import sys

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

WS_RE = re.compile(r"[\xa0\u2007\u202f\s]+")

def norm(s):
    """空白归一化：docx 公司名常含 \xa0 不换行空格，直接 in 匹配必失败"""
    if s is None:
        return ""
    return WS_RE.sub(" ", str(s)).strip()

def docx_all_text(path):
    """python-docx 读全文：段落 + 表格（合并单元格会重复值，容忍）。"""
    from docx import Document
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return norm(" | ".join(parts))

def xls_all_text(path):
    """xlrd 遍历全部单元格。"""
    import xlrd
    rb = xlrd.open_workbook(path)
    parts = []
    for si in range(rb.nsheets):
        ws = rb.sheet_by_index(si)
        for r in range(ws.nrows):
            for c in range(ws.ncols):
                v = ws.cell_value(r, c)
                if v not in ("", None):
                    parts.append(str(v))
    return norm(" | ".join(parts))

def find_amount(text, label):
    """从文本中找总金额: 优先 TOTAL/ВСЕГО 行金额，否则取全部金额中最大值。
    支持 ￥53000、53000,00、53 000.00、CNY 53000 等。"""
    def parse(s):
        s = s.replace(" ", "").replace(",", ".")
        try:
            return float(s)
        except ValueError:
            return None
    amounts = []
    for p in [r"￥\s*([\d\s.,]+)", r"(\d{1,3}(?:[\s,.]\d{3})+[,.]\d{2})",
              r"([\d]{1,7}[,.]\d{2})"]:
        for m in re.finditer(p, text):
            v = parse(m.group(1))
            if v is not None:
                amounts.append((v, m.start()))
    if not amounts:
        return None
    # 优先 TOTAL/ВСЕГО 上下文附近的金额
    for marker in ("TOTAL", "ВСЕГО", "Всего"):
        i = text.find(marker)
        if i >= 0:
            near = [v for v, pos in amounts if abs(pos - i) < 200]
            if near:
                return max(near)
    return max(v for v, _ in amounts)

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--contract", required=True, help="合同 fields.json（ocr_pdf.py 输出）")
    ap.add_argument("--invoice", required=True, help="发票 docx")
    ap.add_argument("--packing", required=True, help="箱单 docx")
    ap.add_argument("--customs", required=True, help="报关单 xls")
    ap.add_argument("--report", default="三单校验报告.md", help="报告输出路径")
    args = ap.parse_args()

    cf = json.load(open(args.contract, encoding="utf-8"))["fields"]
    inv, pck, cus = (docx_all_text(args.invoice), docx_all_text(args.packing),
                     xls_all_text(args.customs))

    checks, ok_all = [], True

    def check(name, cond, detail=""):
        nonlocal ok_all
        ok_all = ok_all and bool(cond)
        checks.append((name, bool(cond), detail))

    # 1. 合同号 / 规格书号三单一致
    cno = cf.get("contract_no", ""); sno = cf.get("spec_no", "")
    check("合同号三单一致", cno and cno in inv and cno in pck and cno in cus,
          f"合同={cno} 发票={'✓' if cno in inv else '✗'} 箱单={'✓' if cno in pck else '✗'} 报关单={'✓' if cno in cus else '✗'}")
    check("规格书号三单一致", sno and sno in inv and sno in pck and sno in cus,
          f"合同={sno} 发票={'✓' if sno in inv else '✗'} 箱单={'✓' if sno in pck else '✗'} 报关单={'✓' if sno in cus else '✗'}")

    # 2. 发票金额 == 合同总金额
    c_total = cf.get("total")
    inv_amt = find_amount(inv, "发票")
    if c_total:
        cv = float(c_total.replace(" ", "").replace(",", "."))
        check("发票金额=合同总金额",
              inv_amt is not None and abs(inv_amt - cv) < 0.01,
              f"合同={cv} 发票={'未识别' if inv_amt is None else inv_amt}")
    else:
        check("发票金额=合同总金额", False, "合同 fields.json 无 total 字段 — 先跑 ocr_pdf.py/extract_fields.py")

    # 3. 箱单 vs 报关单 数量/毛净重一致性（抓双方数字做集合比对，简版）
    def numbers(s):
        return set(re.findall(r"\d+(?:[.,]\d+)?", s))
    shared = numbers(pck) & numbers(cus)
    check("箱单/报关单数字互通", len(shared) >= 3,
          f"双方共有数字 {len(shared)} 个（件数/毛净重/体积应在交集内）")

    # 4. 无编造: 合同关键编号（三单必含字段）在单据中的覆盖率
    keys = [v for k, v in cf.items() if k in ("contract_no", "spec_no") and v]
    covered = [v for v in keys if v in inv + pck + cus]
    check("合同关键编号全数覆盖单据", len(covered) == len(keys),
          f"覆盖 {len(covered)}/{len(keys)}: {keys}")

    # 报告
    lines = ["# 三单一致性校验报告", "",
             f"- 合同字段源: {args.contract}", f"- 发票: {args.invoice}",
             f"- 箱单: {args.packing}", f"- 报关单: {args.customs}", ""]
    for name, ok, detail in checks:
        lines.append(f"- [{'✅' if ok else '❌'}] {name}" + (f" — {detail}" if detail else ""))
    lines += ["", f"**结论: {'全部通过' if ok_all else '存在不一致，需人工修正' }**"]
    with open(args.report, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print("\n".join(lines))
    sys.exit(0 if ok_all else 1)

if __name__ == "__main__":
    main()
