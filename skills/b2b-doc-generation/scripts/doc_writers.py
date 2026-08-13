#!/usr/bin/env python3
"""
doc_writers.py — 从 order.json 生成单据（套真实模板 / 自绘两模式）

套模板模式（--template-dir 指向含真实模板的目录，如 Desktop\\发运文件\\）：
  - 报关单.xls：xlutils.copy 复制真实模板 → 只改值单元格 → 合并单元格/边框 100% 保留
  - *发票*.docx / *箱单*.docx：python-docx 复制模板 → 改字段

自绘模式（无 --template-dir）：从零画近似版式（fallback）

用法：
  python doc_writers.py all <order.json> [--config companies.yaml] [--template-dir 目录] [--outdir 目录]
  python doc_writers.py invoice|packing|customs <order.json> [...]

依赖：python-docx、xlrd、xlwt、xlutils、pyyaml（可选）。
"""
import argparse
import json
import sys
from pathlib import Path

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

from docx import Document
from docx.shared import Pt
import xlwt

try:
    import yaml
    HAS_YAML = True
except ImportError:
    HAS_YAML = False


# ── 数据加载 ──────────────────────────────────────────────────────


def load_order(path):
    return json.loads(Path(path).read_text(encoding="utf-8"))


def load_config(path):
    if not path or not HAS_YAML:
        return {}
    with open(path, "r", encoding="utf-8") as f:
        return yaml.safe_load(f) or {}


def resolve_seller(order, config):
    sid = order.get("seller_id")
    return (config.get("sellers") or {}).get(sid, {}) if sid else {}


def resolve_buyer(order, config):
    bid = order.get("buyer_id")
    return (config.get("buyers") or {}).get(bid, {}) if bid else {}


def find_template(template_dir, *patterns):
    """在目录里按文件名模式找模板，返回第一个匹配的路径。"""
    d = Path(template_dir)
    if not d.is_dir():
        return None
    for p in patterns:
        for f in d.iterdir():
            if f.is_file() and f.name.lower().startswith("~$"):
                continue  # 跳过 Word 锁文件
            if f.is_file() and any(k in f.name for k in p):
                return str(f)
    return None


# ════════════════════════════════════════════════════════════════
# 套模板：报关单（Excel COM 打开改值保存，合并单元格/边框 100% 保留）
# 注：xlutils.copy 会丢边框（纯 copy 实测 354 格边框不一致），故改用 Excel COM。
# ════════════════════════════════════════════════════════════════


def generate_customs_from_template(order, config, template_path, out_path):
    """套报关单模板：用 Excel COM 打开 → 改值 → 另存。

    为什么不用 xlutils.copy：xlutils.copy 复制 .xls 时会丢失边框/字体（纯 copy
    实测 354 格边框不一致）。Excel COM 由 Excel 本体打开/保存，格式 100% 保留。
    前置：Windows + 安装 Excel/WPS（pywin32）。
    """
    import win32com.client

    seller = resolve_seller(order, config)
    terms = order.get("terms", {}) or {}
    items = order.get("items", [])
    packing = order.get("packing", {}) or {}

    country = terms.get("destination_country", "俄罗斯")
    currency_zh = {"CNY": "人民币", "USD": "美元", "EUR": "欧元", "RUB": "卢布"}.get(
        (terms.get("currency") or "USD").upper(), "美元")

    app = win32com.client.DispatchEx("Excel.Application")
    app.Visible = False
    app.DisplayAlerts = False
    try:
        wb = app.Workbooks.Open(template_path)
        ws = wb.Worksheets(1)

        def setv(r, c, v):
            ws.Cells(r + 1, c + 1).Value = v  # Excel 坐标 1 基

        # ── 抬头字段（0 基坐标，按真实报关单.xls 版式）──
        seller_line = f"{seller.get('name_cn', '')}{seller.get('uscc', '')}"
        setv(5, 0, seller_line)                       # 收发货人
        setv(7, 0, seller_line)                       # 生产销售单位
        setv(11, 1, country)                          # 贸易国
        setv(11, 3, country)                          # 运抵国
        setv(11, 10, seller.get("origin_place", "沈阳"))  # 境内货源地
        setv(13, 3, terms.get("incoterm", "FOB"))     # 成交方式
        setv(15, 0, f"{order.get('contract_no', '')} {order.get('order_no', '')}")  # 合同协议号
        setv(15, 3, packing.get("package_count", ""))       # 件数
        setv(15, 4, packing.get("package_type", "木箱"))    # 包装种类
        setv(15, 7, packing.get("gross_weight_kg", ""))     # 毛重
        setv(14, 9, f"净重(千克){packing.get('net_weight_kg', '')}")  # 净重

        # ── 商品区（模板 9 项，R21 起每项 2 行）──
        TEMPLATE_ITEM_ROWS = 9
        first_row = 21
        for i, it in enumerate(items):
            if i >= TEMPLATE_ITEM_ROWS:
                print(f"  ⚠ 商品 {i + 1} 项超出模板容量（{TEMPLATE_ITEM_ROWS} 项），已跳过。"
                      "多于此数需扩展报关单模板或分单。", file=sys.stderr)
                break
            r = first_row + 2 * i
            nw = it.get("total_weight_kg") or it.get("weight_kg_per_unit") or 0
            qty = it.get("quantity", 0)
            unit = it.get("unit", "")
            setv(r, 0, i + 1)                            # 项号
            setv(r, 1, str(it.get("hs_code", "")))       # 商品编号
            setv(r, 2, it.get("description_cn") or it.get("description_en", ""))  # 品名
            setv(r, 4, f"{nw}千克\n{qty}{unit}")         # 数量及单位
            setv(r, 6, country)                          # 最终目的国
            setv(r, 8, it.get("unit_price") or "")       # 单价
            setv(r, 9, it.get("amount") or "")           # 总价
            setv(r, 10, currency_zh)                     # 币制
            # 申报要素行(r+1)：保留模板原文，不覆盖

        # 清空多余商品行
        for i in range(len(items), TEMPLATE_ITEM_ROWS):
            r = first_row + 2 * i
            for c in (0, 1, 2, 4, 6, 8, 9, 10):
                setv(r, c, "")
            setv(r + 1, 0, "")  # 申报要素行

        wb.SaveAs(out_path, FileFormat=56)  # 56 = xlExcel8 (.xls)
        wb.Close(False)
    finally:
        app.Quit()
    return out_path


# ════════════════════════════════════════════════════════════════
# 套模板：发票 / 箱单（python-docx 复制 + 改字段）
# ════════════════════════════════════════════════════════════════


def _set_para_text(para, text):
    """重写段落文本（清空 runs 后写新文本，保留段落格式）。"""
    for r in list(para.runs):
        r.text = ""
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)


def _set_cell_text(cell, text):
    """重写单元格文本（保留首段格式，其余段删除）。"""
    p = cell.paragraphs[0]
    for r in list(p.runs):
        r.text = ""
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def generate_invoice_from_template(order, config, template_path, out_path):
    seller = resolve_seller(order, config)
    buyer = resolve_buyer(order, config)
    terms = order.get("terms", {}) or {}
    items = order.get("items", [])
    cur = terms.get("currency", "USD")
    sym = {"CNY": "￥", "USD": "$", "EUR": "€", "RUB": "₽"}.get(cur.upper(), "$")

    doc = Document(template_path)
    t = doc.tables[0]
    buyer_name = buyer.get("name_en", "") or buyer.get("name_ru", "")

    # 抬头字段（按 26LF031发票.docx 版式坐标）
    _set_cell_text(t.cell(0, 0), f"Seller\n{seller.get('name_en', '')}")
    _set_cell_text(t.cell(2, 0), f"Buyer\n{buyer_name}")
    _set_cell_text(t.cell(1, 2), f"Invoice\n{order.get('order_no', '')}")
    _set_cell_text(t.cell(3, 2), f"Contract /Договор No.\n{order.get('contract_no', '')}")
    _set_cell_text(t.cell(3, 4), f"Date /Дата\n{order.get('date', '')}")
    port = terms.get("port_of_loading", "") or terms.get("destination", "")
    _set_cell_text(t.cell(4, 0), f"Terms of delivery \n{terms.get('incoterm', '')} Port of {port}")
    _set_cell_text(t.cell(5, 2), f"Terms of Payment: {terms.get('payment', 'T/T')}")

    # 商品行（模板 9 项，R7 起）
    first = 7
    TEMPLATE_ITEM_ROWS = 9
    for i, it in enumerate(items):
        if i >= TEMPLATE_ITEM_ROWS:
            break
        r = first + i
        desc = it.get("description_ru") or it.get("description_en") or it.get("description_cn")
        qty = it.get("quantity", 0)
        up = it.get("unit_price")
        am = it.get("amount")
        _set_cell_text(t.cell(r, 0), str(i + 1))
        _set_cell_text(t.cell(r, 1), desc)
        _set_cell_text(t.cell(r, 2), f"{qty:,.0f}")
        _set_cell_text(t.cell(r, 3), f"{up:,.2f}" if up is not None else "")
        _set_cell_text(t.cell(r, 4), f"{am:,.2f}" if am is not None else "")
    # 清空多余商品行
    for i in range(len(items), TEMPLATE_ITEM_ROWS):
        r = first + i
        for c in range(5):
            _set_cell_text(t.cell(r, c), "")

    # Total 行
    total_qty = sum(it.get("quantity", 0) for it in items)
    total = sum(it.get("amount") or 0 for it in items)
    _set_cell_text(t.cell(16, 1), "Total")
    _set_cell_text(t.cell(16, 2), f"{total_qty:,.0f}")
    _set_cell_text(t.cell(16, 4), f"＄{total:,.0f}")

    doc.save(out_path)
    return out_path


def generate_packing_from_template(order, config, template_path, out_path):
    seller = resolve_seller(order, config)
    buyer = resolve_buyer(order, config)
    terms = order.get("terms", {}) or {}
    items = order.get("items", [])

    doc = Document(template_path)
    buyer_name = buyer.get("name_en", "") or buyer.get("name_ru", "")

    # 段落字段（按 26LF031箱单.docx 版式：公司名→买家→合同号→条款）
    paras = doc.paragraphs
    port = terms.get("port_of_loading", "") or terms.get("destination", "")
    for p in paras:
        txt = p.text
        if "Packing List" in txt or "Упаковочный" in txt or "TO MESSRS" in txt or "К ГОСПОДАМ" in txt:
            continue  # 标题/抬头标签不动
        if seller.get("name_en") and seller["name_en"].split()[0] in txt:
            _set_para_text(p, seller.get("name_en", ""))
        elif "LIMITED LIABILITY" in txt or "INTERPIPE" in txt.upper() or "ООО" in txt.upper():
            _set_para_text(p, buyer_name)
        elif "Contract No." in txt or "Contract" in txt:
            _set_para_text(p, f"Contract No. {order.get('contract_no', '')}                                                  Date:  {order.get('date', '')}")
        elif "Terms of delivery" in txt:
            _set_para_text(p, f"Terms of delivery:  {terms.get('incoterm', '')} Port of  {port}                    Country of origin:  China")
    # 落款公司名（倒数第二段）
    if len(paras) >= 2 and "Terms of delivery" not in paras[-2].text and "Contract" not in paras[-2].text:
        _set_para_text(paras[-2], seller.get("name_en", ""))

    t = doc.tables[0]
    first = 1
    TEMPLATE_ITEM_ROWS = 9
    for i, it in enumerate(items):
        if i >= TEMPLATE_ITEM_ROWS:
            break
        r = first + i
        desc = it.get("description_ru") or it.get("description_en") or it.get("description_cn")
        _set_cell_text(t.cell(r, 0), desc)
        _set_cell_text(t.cell(r, 1), f"{it.get('quantity', 0):,.0f}")
        _set_cell_text(t.cell(r, 2), it.get("unit", "pcs"))
    for i in range(len(items), TEMPLATE_ITEM_ROWS):
        r = first + i
        for c in range(7):
            _set_cell_text(t.cell(r, c), "")

    total_qty = sum(it.get("quantity", 0) for it in items)
    _set_cell_text(t.cell(10, 0), "Total")
    _set_cell_text(t.cell(10, 1), f"{total_qty:,.0f}")

    doc.save(out_path)
    return out_path


# ════════════════════════════════════════════════════════════════
# 自绘 fallback（无模板时）
# ════════════════════════════════════════════════════════════════


def currency_symbol(cur):
    return {"CNY": "￥", "USD": "$", "EUR": "€", "RUB": "₽"}.get((cur or "USD").upper(), (cur or "USD"))


def generate_invoice(order, config, out_path):
    seller = resolve_seller(order, config)
    buyer = resolve_buyer(order, config)
    terms = order.get("terms", {}) or {}
    cur = terms.get("currency", "USD")
    sym = currency_symbol(cur)
    doc = Document()
    table = doc.add_table(rows=0, cols=5)
    table.style = "Table Grid"

    def add_merged(text, bold=False, size=10):
        row = table.add_row().cells
        merged = row[0]
        for c in row[1:]:
            merged = merged.merge(c)
        r = merged.paragraphs[0].add_run(text)
        r.bold = bold
        r.font.size = Pt(size)

    add_merged(f"Suppler/ Поставщик/ {seller.get('name_en', '')}", bold=True)
    if seller.get("address"):
        add_merged(seller.get("address"), size=9)
    buyer_line = buyer.get("name_ru", "") or buyer.get("name_en", "")
    if buyer.get("name_en") and buyer.get("name_ru"):
        buyer_line = f"{buyer.get('name_ru')} / {buyer.get('name_en')}"
    add_merged(f"Buyer/ Покупатель/ {buyer_line}", bold=True)
    if buyer.get("address"):
        add_merged(buyer.get("address"), size=9)
    dest = terms.get("destination", "") or terms.get("port_of_destination", "")
    add_merged(f"Terms of delivery /Условия поставки/ {terms.get('incoterm', '')} {dest}   |   "
               f"Country of origin: China /Страна происхождения: Китай", size=9)
    add_merged(f"Terms of Payment: {terms.get('payment', 'T/T')} /Условия оплаты: Т/Т", size=9)

    headers = ["Marks & numbers//\nНаименование и номера",
               "Number & kind of packages, description of goods//\nКоличество и вид упаковок, описание товара",
               "Quantity//\nКоличество", "Unit price//\nЦена", "Amount//\nСумма"]
    hrow = table.add_row().cells
    for i, h in enumerate(headers):
        r = hrow[i].paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(8)

    for it in order.get("items", []):
        desc = it.get("description_ru") or it.get("description_en") or it.get("description_cn")
        if it.get("standard"):
            desc = f"{desc} {it.get('standard')}"
        cells = table.add_row().cells
        cells[0].text = it.get("hs_code", "")
        cells[1].text = desc
        cells[2].text = f"{it.get('quantity')} {it.get('unit', '')}".strip()
        up, am = it.get("unit_price"), it.get("amount")
        cells[3].text = f"{sym}{up:,.2f}" if up is not None else ""
        cells[4].text = f"{sym}{am:,.2f}" if am is not None else ""

    total = sum(it.get("amount") or 0 for it in order.get("items", []))
    trow = table.add_row().cells
    t0 = trow[0]
    for c in trow[1:4]:
        t0 = t0.merge(c)
    r = t0.paragraphs[0].add_run(f"TOTAL /ВСЕГО: {sym}{total:,.2f}")
    r.bold = True
    r.font.size = Pt(10)
    trow[4].text = f"{sym}{total:,.2f}"
    add_merged(seller.get("name_en", ""), bold=True)
    add_merged(order.get("date", ""), size=9)
    doc.save(out_path)
    return out_path


def generate_packing_list(order, config, out_path):
    seller = resolve_seller(order, config)
    buyer = resolve_buyer(order, config)
    terms = order.get("terms", {}) or {}
    dest = terms.get("destination", "") or terms.get("port_of_destination", "")
    doc = Document()
    buyer_ru = buyer.get("name_ru", "") or buyer.get("name_en", "")
    buyer_en = buyer.get("name_en", "")
    paragraphs = [
        seller.get("name_en", ""), seller.get("address", ""), "Packing List", "Упаковочный лист",
        "TO MESSRS/ К ГОСПОДАМ:",
        f"{buyer_ru} / {buyer_en}" if buyer_en and buyer_ru != buyer_en else buyer_ru,
        buyer.get("address", ""),
        f"Contract/ Договор No. {order.get('contract_no', '')}    Date /Дата: {order.get('date', '')}",
        f"Specification /Спецификация No. {order.get('spec_no', '')}",
        f"Terms of delivery:  {terms.get('incoterm', '')} {dest}    Country of origin:  China",
        f"Условия поставки: {terms.get('incoterm', '')} {dest}    Страна происхождения: Китай",
        seller.get("name_en", ""), order.get("date", ""),
    ]
    for txt in paragraphs:
        p = doc.add_paragraph()
        r = p.add_run(txt if txt else "")
        r.font.size = Pt(10)
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ["Description of goods//Описание товаров", "Qty//Кол-во", "Unit//Единица измерения",
               "Package//Упаковка", "N.W.(kg)//Вес нетто", "G.W.(kg)//Вес брутто", "Volume/m3//Объём/м3"]
    for i, h in enumerate(headers):
        r = table.rows[0].cells[i].paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(8)
    for it in order.get("items", []):
        desc = it.get("description_ru") or it.get("description_en") or it.get("description_cn")
        nw = it.get("total_weight_kg") or it.get("weight_kg_per_unit")
        gw = round((nw or 0) * 1.05, 2) if nw else ""
        cells = table.add_row().cells
        for j, v in enumerate([desc, it.get("quantity", 0), it.get("unit", ""), "", nw or "", gw or "", ""]):
            cells[j].text = str(v)
    total_qty = sum(it.get("quantity", 0) for it in order.get("items", []))
    total_nw = sum(it.get("total_weight_kg") or it.get("weight_kg_per_unit") or 0 for it in order.get("items", []))
    cells = table.add_row().cells
    for j, v in enumerate(["Total//Всего", total_qty, "", "", round(total_nw, 2) or "",
                           round(total_nw * 1.05, 2) if total_nw else "", ""]):
        cells[j].text = str(v)
    doc.save(out_path)
    return out_path


def generate_customs(order, config, out_path):
    seller = resolve_seller(order, config)
    terms = order.get("terms", {}) or {}
    items = order.get("items", [])
    wb = xlwt.Workbook(encoding="utf-8")
    ws = wb.add_sheet("出口报关单", cell_overwrite_ok=True)
    title = xlwt.easyxf("font: bold on, height 280; align: horiz center, vert centre;")
    cell = xlwt.easyxf("font: height 200; align: horiz left, vert centre, wrap on;"
                       "borders: left thin, right thin, top thin, bottom thin;")
    lbl = xlwt.easyxf("font: height 200, bold on; align: horiz left, vert centre;"
                      "borders: left thin, right thin, top thin, bottom thin;")

    def w(r, c, v, style=cell):
        ws.write(r, c, v, style)

    w(0, 0, "中华人民共和国海关出口货物报关单", title)
    w(4, 0, "收发货人", lbl); w(5, 0, f"{seller.get('name_cn', '')}{seller.get('uscc', '')}")
    w(10, 0, "贸易国(地区)", lbl); w(11, 1, "俄罗斯")
    w(12, 0, "成交方式", lbl); w(13, 3, terms.get("incoterm", "DAP"))
    w(14, 0, "合同协议号", lbl); w(15, 0, f"{order.get('contract_no', '')}/{order.get('spec_no', '')}")
    heads = ["项号", "商品编号", "商品名称、规格型号", "数量及单位", "最终目的国(地区)", "单价", "总价", "币制", "征免"]
    for c, h in enumerate(heads):
        w(20, c, h, lbl)
    r = 21
    for i, it in enumerate(items, 1):
        w(r, 0, i); w(r, 1, it.get("hs_code", ""))
        w(r, 2, it.get("description_cn") or it.get("description_en", ""))
        w(r, 3, f"{it.get('quantity')}{it.get('unit', '')}")
        w(r, 4, "俄罗斯"); w(r, 5, it.get("unit_price") or ""); w(r, 6, it.get("amount") or "")
        w(r, 7, "人民币" if terms.get("currency", "USD") == "CNY" else "美元"); w(r, 8, "照章")
        r += 1
    wb.save(out_path)
    return out_path


# ── CLI ───────────────────────────────────────────────────────────


def main(argv):
    ap = argparse.ArgumentParser(description="从 order.json 生成单据（套真实模板 / 自绘）")
    ap.add_argument("mode", choices=["all", "invoice", "packing", "customs"])
    ap.add_argument("order", help="order.json 路径")
    ap.add_argument("--config", default=None, help="companies.yaml 路径")
    ap.add_argument("--template-dir", default=None, help="真实模板目录（含报关单.xls/*发票*.docx/*箱单*.docx）")
    ap.add_argument("--outdir", default=None, help="输出目录（默认 order.json 同目录）")
    args = ap.parse_args(argv)

    order = load_order(args.order)
    config = load_config(args.config)
    outdir = Path(args.outdir) if args.outdir else Path(args.order).parent
    outdir.mkdir(parents=True, exist_ok=True)
    order_no = order.get("order_no", Path(args.order).stem)
    td = args.template_dir

    def out(ext):
        return str(outdir / f"{order_no}_{ext}")

    modes = ["invoice", "packing", "customs"] if args.mode == "all" else [args.mode]
    for m in modes:
        p = None
        if m == "customs":
            tpl = find_template(td, ("报关单", "customs")) if td else None
            if tpl:
                p = generate_customs_from_template(order, config, tpl, out("customs.xls"))
                print(f"✅ 已生成 customs（套模板）: {p}")
            else:
                p = generate_customs(order, config, out("customs.xls"))
                print(f"✅ 已生成 customs（自绘近似版）: {p}")
        elif m == "invoice":
            tpl = find_template(td, ("发票", "invoice")) if td else None
            if tpl:
                p = generate_invoice_from_template(order, config, tpl, out("invoice.docx"))
                print(f"✅ 已生成 invoice（套模板）: {p}")
            else:
                p = generate_invoice(order, config, out("invoice.docx"))
                print(f"✅ 已生成 invoice（自绘）: {p}")
        else:
            tpl = find_template(td, ("箱单", "packing")) if td else None
            if tpl:
                p = generate_packing_from_template(order, config, tpl, out("packing.docx"))
                print(f"✅ 已生成 packing（套模板）: {p}")
            else:
                p = generate_packing_list(order, config, out("packing.docx"))
                print(f"✅ 已生成 packing（自绘）: {p}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
