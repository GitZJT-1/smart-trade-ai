#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
dump_template.py — 模板结构自动 dump（docx / xls）

首次遇到新模板跑一遍 → 生成结构档案（段落顺序、表格行列、表头文本、
非空单元格坐标），存档进 references/ 即可替代人工记录 sample-formats.md。

用法:
  python dump_template.py <模板.docx 或 模板.xls> [--out struct.md]
  python dump_template.py <模板.xls> --cells    # 附带非空单元格坐标

输出 markdown: 段落顺序 / 表格(行×列+表头) / 非空单元格坐标
"""
import argparse
import sys

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

def dump_docx(path):
    from docx import Document
    doc = Document(path)
    lines = [f"# 模板结构: {path}", ""]
    lines.append(f"**段落 {len(doc.paragraphs)} 段，表格 {len(doc.tables)} 个**")
    lines.append("")
    lines.append("## 段落顺序")
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t:
            lines.append(f"P{i}: {t}")
        else:
            lines.append(f"P{i}: (空)")
    for ti, t in enumerate(doc.tables):
        lines.append("")
        lines.append(f"## 表格 {ti}: {len(t.rows)} 行 × {len(t.columns)} 列")
        for ri, row in enumerate(t.rows):
            cells = [c.text.replace("\n", "⏎").strip() for c in row.cells]
            # 去重相邻重复（合并单元格 python-docx 会重复值）
            dedup = []
            for c in cells:
                if not dedup or dedup[-1] != c:
                    dedup.append(c)
            lines.append(f"R{ri}: " + " | ".join(dedup))
    return "\n".join(lines)

def dump_xls(path, with_cells=False):
    import xls_tpl
    import xlrd
    rb = xlrd.open_workbook(path, formatting_info=True)
    lines = [f"# 模板结构: {path}", ""]
    for si in range(rb.nsheets):
        ws = rb.sheet_by_index(si)
        lines.append(f"## Sheet {si}「{ws.name}」: {ws.nrows} 行 × {ws.ncols} 列")
        for r in range(ws.nrows):
            vals = [ws.cell_value(r, c) for c in range(ws.ncols)]
            if any(v not in ("", None) for v in vals):
                lines.append(f"R{r}: " + " | ".join("" if v in ("", None) else str(v).replace("\n", "⏎") for v in vals))
        if with_cells:
            lines.append("")
            lines.append("### 非空单元格坐标")
            lines.append("```")
            lines.append(xls_tpl.dump_cells_human(path, si))
            lines.append("```")
    return "\n".join(lines)

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("file", help="模板 docx 或 xls")
    ap.add_argument("--out", help="输出 md 路径（默认 stdout）")
    ap.add_argument("--cells", action="store_true", help="xls 附带非空坐标")
    args = ap.parse_args()

    low = args.file.lower()
    if low.endswith(".docx"):
        text = dump_docx(args.file)
    elif low.endswith(".xls"):
        text = dump_xls(args.file, args.cells)
    else:
        sys.exit(f"不支持类型: {low}（支持 .docx / .xls）")
    if args.out:
        with open(args.out, "w", encoding="utf-8") as f:
            f.write(text)
        print(f"[OK] → {args.out}")
    else:
        print(text)

if __name__ == "__main__":
    main()
